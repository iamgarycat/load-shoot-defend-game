from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "game_theory_science_fair_report.docx"
DATA = json.loads((ROOT / "report_data.json").read_text(encoding="utf-8"))

EAST_ASIA_FONT = "DFKai-SB"
LATIN_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"
CODE_FONT = "Consolas"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9D9D9"
CONTENT_DXA = 9638


def set_run_font(
    run,
    size=12,
    bold=None,
    italic=None,
    color=BLACK,
    east_asia=EAST_ASIA_FONT,
    latin=LATIN_FONT,
):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-TW")


def set_style_font(style, size, bold=False, color=BLACK):
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr, value in (
        ("ascii", LATIN_FONT),
        ("hAnsi", LATIN_FONT),
        ("cs", LATIN_FONT),
        ("eastAsia", EAST_ASIA_FONT),
    ):
        rfonts.set(qn(f"w:{attr}"), value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="808080", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def format_table_text(table, size=9.6, centered_cols=()):
    for ridx, row in enumerate(table.rows):
        for cidx, cell in enumerate(row.cells):
            if ridx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if ridx == 0 or cidx in centered_cols:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run, size=size, bold=(ridx == 0))


def set_paragraph_shading(paragraph, fill=LIGHT_GRAY):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def add_body(doc, text, *, first_indent=True, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(24) if first_indent else Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text.strip())
    set_run_font(run, italic=italic)
    return p


def add_paragraphs(doc, text):
    for paragraph in [p.strip() for p in text.strip().split("\n\n") if p.strip()]:
        add_body(doc, paragraph)


def add_labeled(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(label)
    set_run_font(r1, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_note(doc, label, text, fill=LIGHT_GRAY):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, fill)
    r1 = p.add_run(f"{label}　")
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_numbered(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 14, 3: 12}[level], bold=True)
    return p


def add_chapter(doc, title, *, first=False):
    return add_heading(doc, title, 1)


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, size=11.5, italic=True, east_asia=EAST_ASIA_FONT, latin=MATH_FONT)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=GRAY)
    return p


def add_figure(doc, path, caption, alt, width_cm=15.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    for doc_pr in run._r.xpath(".//wp:docPr"):
        doc_pr.set("descr", alt)
        doc_pr.set("title", caption)
    add_caption(doc, caption)


def add_code(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph(style="Code")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_paragraph_shading(p, "F7F7F7")
        r = p.add_run(line or " ")
        set_run_font(r, size=8.3, east_asia=EAST_ASIA_FONT, latin=CODE_FONT)


def make_table(doc, headers, rows, widths, *, size=9.6, centered_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = str(value)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    set_table_geometry(table, widths)
    format_table_text(table, size=size, centered_cols=centered_cols)
    return table


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=10)


def set_page_number_format(section, fmt, start=1):
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:pgNumType"))
    if old is not None:
        sect_pr.remove(old)
    node = OxmlElement("w:pgNumType")
    node.set(qn("w:fmt"), fmt)
    node.set(qn("w:start"), str(start))
    sect_pr.append(node)


def setup_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)


def set_footer_page_number(section, fmt, start=1):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    for run in list(p.runs):
        p._p.remove(run._r)
    add_page_field(p)
    set_page_number_format(section, fmt, start)


def add_toc_entry(doc, text, page, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Cm(0 if level == 1 else 0.7)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5 if level == 1 else 1)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(text)
    set_run_font(r1, size=12, bold=(level == 1))
    r2 = p.add_run(f"\t{page}")
    set_run_font(r2, size=12, bold=(level == 1))


def configure_document():
    doc = Document()
    doc.core_properties.title = "裝彈、射擊與防禦遊戲的無限期零和賽局分析"
    doc.core_properties.subject = "有限截斷的收斂、無限期價值與第二回合均衡"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "高中科展, 賽局理論, 納許均衡, 無限期賽局, 電腦輔助證明"

    styles = doc.styles
    set_style_font(styles["Normal"], 12, False)
    normal_pf = styles["Normal"].paragraph_format
    normal_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_pf.first_line_indent = Pt(24)
    normal_pf.space_before = Pt(0)
    normal_pf.space_after = Pt(0)
    normal_pf.line_spacing = 1.0

    for name, size, before, after in (
        ("Heading 1", 16, 12, 6),
        ("Heading 2", 14, 8, 4),
        ("Heading 3", 12, 6, 2),
    ):
        set_style_font(styles[name], size, True)
        pf = styles[name].paragraph_format
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.0
        pf.keep_with_next = True

    for list_name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        set_style_font(styles[list_name], 12, False)
        pf = styles[list_name].paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

    set_style_font(styles["Caption"], 10.5, False, GRAY)
    styles["Caption"].paragraph_format.space_before = Pt(2)
    styles["Caption"].paragraph_format.space_after = Pt(4)
    styles["Caption"].paragraph_format.first_line_indent = Pt(0)

    if "Code" not in styles:
        code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code"]
    set_style_font(code_style, 8.3, False)

    setup_section(doc.sections[0])
    doc.sections[0].footer.is_linked_to_previous = False
    doc.sections[0].footer.paragraphs[0].text = ""
    return doc


doc = configure_document()

# Cover
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(26)
r = p.add_run("高中科學展覽成果報告")
set_run_font(r, size=16, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run("裝彈、射擊與防禦遊戲的無限期零和賽局分析")
set_run_font(r, size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(48)
r = p.add_run("有限截斷的收斂、無限期價值與第二回合均衡")
set_run_font(r, size=14)

for line in (
    "科　　別：請填寫",
    "學校名稱：請填寫",
    "作者姓名：請填寫",
    "就讀年級：請填寫",
    "指導老師：請填寫",
    "完成日期：請填寫",
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(line)
    set_run_font(r, size=12)

# TOC section
toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
setup_section(toc_section)
set_footer_page_number(toc_section, "upperRoman", 1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("目　錄")
set_run_font(r, size=16, bold=True)

TOC_PAGES = {
    "三、前言": 1,
    "四、研究目的": 2,
    "五、文獻探討": 3,
    "六、研究方法": 5,
    "七、結果與討論（含結論與建議）": 16,
    "八、參考文獻": 21,
    "九、附錄": 21,
}
for title, page in TOC_PAGES.items():
    add_toc_entry(doc, title, page)
add_note(doc, "頁碼說明", "目錄頁使用大寫羅馬數字；正文至附錄使用阿拉伯數字並連續編號。本目錄列至章級，章內小節可由 Word 的導覽窗格查看。")

# Body section
body_section = doc.add_section(WD_SECTION.NEW_PAGE)
setup_section(body_section)
set_footer_page_number(body_section, "decimal", 1)

# CONTENT_START
add_chapter(doc, "三、前言", first=True)

add_heading(doc, "3.1　摘要", 2)
add_paragraphs(
    doc,
    """
本研究分析一個雙人同時行動的零和遊戲。每回合可選擇裝彈、射擊或防禦；射擊需要消耗一發子彈，防禦可以擋住射擊但不能連續使用。若一方射擊，而另一方既未射擊也未防禦，射擊者立即獲勝；雙方同時射擊則各消耗一發子彈並繼續。玩家可以不斷裝彈，所以遊戲有無限多個狀態，也沒有預先固定的結束回合。

本研究先把遊戲截斷為最多 N 回合，利用有限零和矩陣賽局從後往前計算。接著分別把截止時尚未結束的局面記為 -1 與 +1，得到悲觀值 V̲_N 與樂觀值 V̄_N。本文證明 V̲_N 單調上升、V̄_N 單調下降，兩者都有極限。真正的困難是證明兩個極限相等；為此，本研究建立交叉最優策略、間距不等式與資源分數 Φ，並把無限多個狀態分成可手算處理的高資源區，以及只有 37 個有序狀態的低資源核心。低資源核心使用分數運算逐一驗證，不使用浮點數判斷正負。

主要結論是：對每個合法狀態，上下截止極限相等，因此平手截止的有限期值也收斂，且此共同極限就是無限期遊戲的值。若雙方第一回合都裝彈，第二回合狀態為（3,3）；本文進一步證明該狀態的有限期均衡策略會收斂。浮點探索估計裝彈、射擊、防禦機率約為（0.2791962356，0.3076983951，0.4131053693），但此十位小數僅屬數值估計；目前以精確分數嚴格證明的範圍較寬，兩者在報告中分開呈現。
""",
)
add_body(doc, "關鍵詞：賽局理論、二人零和賽局、納許均衡、動態規劃、無限期遊戲、電腦輔助證明。", first_indent=False)

add_heading(doc, "3.2　研究緣起", 2)
add_paragraphs(
    doc,
    """
這個遊戲的規則像猜拳：兩位玩家同時選擇行動，而且單回合的勝負取決於雙方的組合。然而它又比猜拳複雜，因為子彈會累積、防禦有冷卻限制，今天的選擇會改變明天可選的行動。若只看目前一回合，容易忽略長期資源；若只靠固定套路，又會被對方針對。

最直接的計算方法，是先規定遊戲最多進行 N 回合，若仍未分勝負便算平手，再由最後一回合倒推。實驗很快顯示，若 N 增加，關鍵狀態的數值與策略似乎穩定。可是「小數看起來不動」不是數學證明：數列可能緩慢移動，也可能在極小範圍內振盪。科展研究因此從「算很大的 N」轉為「建立可證明的上下界」。

一開始曾猜測關鍵間距可能滿足 G_{N+5}≤0.31G_N，因為數值上每增加五回合，間距大約乘上 0.302 至 0.304。這個猜測若成立，會直接給出指數收斂；但要對所有 N 與所有相關狀態證明同一個常數並不容易。本研究最後採用較穩健的方向：不要求先證明固定收斂速度，只證明上下界的差必定趨近 0。
""",
)

add_heading(doc, "3.3　研究難點", 2)
for item in (
    "狀態數無限：玩家可以一直裝彈，不能只列出一張有限狀態表。",
    "同時行動：每個狀態都要解一個混合策略矩陣賽局，不能只挑單一最佳動作。",
    "有限期與無限期不同：有限 N 的倒推正確，不代表 N→∞ 時一定有唯一極限。",
    "價值收斂與策略收斂不同：即使賽局值收斂，若最優策略不唯一，策略仍可能跳動。",
    "數值與證明不同：浮點數很適合找規律，但嚴格的正負判斷必須控制捨入誤差。",
):
    add_bullet(doc, item)

add_heading(doc, "3.4　研究主線與成果層級", 2)
make_table(
    doc,
    ["層級", "內容", "本文如何處理"],
    [
        ["數學模型", "規則、狀態、行動與轉移", "以整數狀態碼完整表示"],
        ["有限期結果", "N 回合值與混合策略", "Bellman 倒推與有限矩陣賽局"],
        ["無限期定理", "上下截止極限相等", "解析證明＋37 狀態精確證書"],
        ["第二回合策略", "（3,3）策略公式與收斂", "三乘三反對稱矩陣的代數推導"],
        ["小數結果", "近似值與誤差範圍", "精確粗界與浮點探索分開標示"],
    ],
    [1500, 3300, 4838],
    centered_cols=(0,),
)
add_note(doc, "閱讀方式", "第六章是完整證明主體。每個較新的名詞都先翻成有限加權平均、不等式或小矩陣，再說明它在證明中的用途；讀者不必先修高等機率或泛函分析。")


add_chapter(doc, "四、研究目的")
add_heading(doc, "4.1　主要研究問題", 2)
for item in (
    "如何以有限個整數完整編碼子彈數與防禦是否可用？",
    "剩餘 N 回合時，如何計算每個狀態的零和均衡值與混合策略？",
    "當 N→∞ 時，悲觀截止值、平手截止值與樂觀截止值是否收斂到同一數？",
    "如何在狀態數無限的情況下，證明上下極限的間距為 0？",
    "若雙方第一回合都裝彈，第二回合三個行動的有限期均衡策略是否收斂？",
    "哪些小數位是嚴格證明，哪些只是數值觀察？",
):
    add_bullet(doc, item)

add_heading(doc, "4.2　研究目標", 2)
for item in (
    "建立可重現的有限期動態規劃演算法。",
    "用有界單調數列證明上下截止極限存在。",
    "證明交叉最優策略下，當前間距不大於下一回合未終止間距的平均。",
    "設計資源分數，證明交叉過程不會永遠逃向高資源。",
    "把剩餘低資源問題縮成有限個分數不等式並精確驗證。",
    "推出無限期遊戲有值，以及（3,3）有限期策略的極限公式。",
):
    add_numbered(doc, item)

add_heading(doc, "4.3　研究範圍", 2)
add_paragraphs(
    doc,
    """
本文的無限期收益定義為：甲最終獲勝記 +1，乙最終獲勝記 -1；若永遠沒有分出勝負，記 0。本文證明的是二人零和意義下的遊戲值，以及任意 ε>0 都可達到的 ε-最優保證策略。

本文不在尚未證明前宣稱「把每個狀態的極限動作直接拼起來，就一定得到全域精確的平穩納許均衡」。這是比遊戲有值更強的命題。對（3,3）狀態，本文能證明有限期第一步策略收斂，並求得極限矩陣的唯一混合策略；其他所有狀態的策略是否皆唯一，留待後續研究。
""",
)

add_heading(doc, "4.4　研究方法概覽", 2)
make_table(
    doc,
    ["階段", "輸入", "方法", "輸出"],
    [
        ["一", "遊戲規則", "狀態編碼與轉移表", "有限矩陣模型"],
        ["二", "截止收益 z", "倒推與極小極大", "V_N^z"],
        ["三", "z=-1,+1", "單調性與夾擠", "上下極限 V̲、V̄"],
        ["四", "上極限最優策略", "交叉策略與資源分數", "幾乎必定停止"],
        ["五", "停止性", "間距不等式疊代", "V̲=V̄"],
        ["六", "關鍵續局值 A、B", "三乘三矩陣代數", "第二回合極限策略"],
    ],
    [900, 2100, 3000, 3638],
    centered_cols=(0,),
)


add_chapter(doc, "五、文獻探討")
add_heading(doc, "5.1　二人零和矩陣賽局", 2)
add_paragraphs(
    doc,
    """
二人零和表示一方所得等於另一方所失。把甲的純行動列在矩陣的列、乙的純行動列在欄，元素 M_{αβ} 就是雙方選擇 α、β 時甲的收益。混合策略 p 是一組非負機率，總和為 1；若乙選定某一欄 β，甲的平均收益為 Σ_α p_αM_{αβ}。

von Neumann 的極小極大定理指出，在有限二人零和矩陣中，甲先選機率後能保證的最大值，等於乙能把甲壓到的最小值[1]。本文把這個共同數記為 val(M)。這一定理確保每個有限期倒推步驟都有值與最優混合策略；基礎教材可參考[4][5]。為避免把定理當黑箱，6.4 節另證明一個本研究實際使用的「上下保證證書」：只要找到 p、q 同時給出相同上下界，就能用加權平均直接確認它們是均衡。
""",
)

add_heading(doc, "5.2　納許均衡", 2)
add_paragraphs(
    doc,
    """
Nash 於 1950 年提出一般有限賽局的均衡概念[2]。若其他人的策略固定，每位玩家都無法靠單獨改變自己的策略提高收益，這組策略便是納許均衡（Nash equilibrium，亦譯納什均衡）。在二人零和賽局中，甲的極大極小策略與乙的極小極大策略合在一起就是一個納許均衡。

本題不是單一矩陣：每個矩陣元素可能是下一狀態的賽局值。因此有限期策略由剩餘回合數決定，而無限期分析必須同時控制所有後續狀態。
""",
)

add_heading(doc, "5.3　動態規劃與 Bellman 遞迴", 2)
add_paragraphs(
    doc,
    """
動態規劃把長期問題拆成「現在做什麼」與「下一狀態值多少」。本文所稱 Bellman 算子 T，只是一個明確計算程序：先把給定的未來價值函數 W 填入每個一回合矩陣，再取矩陣賽局值。它不是神祕的高等算子；在每個固定狀態，只涉及最多三乘三的有限矩陣。

Shapley 於 1953 年建立隨機賽局的系統模型[3]。一般理論常加入折扣或假設狀態有限；本題沒有折扣，而且子彈數無上限。因此本文不直接套用一般收斂結論，而利用本遊戲的奇偶結構與必勝狀態自行證明。
""",
)

add_heading(doc, "5.4　精確算術與電腦輔助證明", 2)
add_paragraphs(
    doc,
    """
浮點數用有限位元近似實數，連續多次運算後可能產生微小誤差。它適合找規律、畫圖與檢查猜想，但若最終要判斷某個量是否小於或等於 0，就不能只看螢幕上顯示的 0.000000。

Python 的 Fraction 以整數分子與分母保存有理數，加、減、乘、除後仍是精確分數[6]；浮點數的表示限制可參考[7]。本文的低資源證書把每個矩陣元素、聯立方程與不等式都保留為 Fraction；因此「bad=0」代表沒有任何分數真正大於 0，而不是因為四捨五入後看起來等於 0。
""",
)

add_heading(doc, "5.5　本文相對於文獻的工作", 2)
make_table(
    doc,
    ["既有觀念", "本文用途", "本文新增的本題結構"],
    [
        ["極小極大定理", "求每個有限矩陣的值", "矩陣元素由子彈與防禦狀態決定"],
        ["納許均衡", "解釋雙方互為最佳回應", "研究有限期均衡隨 N 的極限"],
        ["動態規劃", "由 N-1 層計算 N 層", "建立上下截止夾擠"],
        ["隨機過程", "追蹤交叉策略造成的狀態", "設計離散資源分數 Φ"],
        ["精確算術", "驗證有限多個不等式", "把無限狀態縮到 37 個核心狀態"],
    ],
    [2200, 3200, 4238],
)

# CHAPTER6_START
add_chapter(doc, "六、研究方法")

add_heading(doc, "6.1　規則、收益與合法行動", 2)
add_paragraphs(
    doc,
    """
固定從玩家甲的角度記收益。甲獲勝為 +1，乙獲勝為 -1，永遠沒有分出勝負為 0。兩人每回合同時選擇行動，不能先看到對方選擇再反應。非法行動不放入當回合矩陣。

裝彈 L 永遠合法，使子彈數增加一。射擊 S 只在至少有一發子彈時合法，並消耗一發。防禦 D 只在上一回合沒有使用防禦時合法；使用後下一回合不能再次防禦。若一方 S、另一方 L，射擊者立即獲勝；若 S/S，雙方各消耗一發後繼續；若 S/D，子彈被擋住並繼續。其他沒有立即命中的組合也繼續。
""",
)
make_table(
    doc,
    ["行動", "合法條件", "資源改變", "對方射擊時"],
    [
        ["裝彈 L", "永遠合法", "子彈 +1；下回合可防禦", "立即落敗"],
        ["射擊 S", "子彈至少 1 發", "子彈 -1；下回合可防禦", "雙方射擊則繼續"],
        ["防禦 D", "本回合可防禦", "子彈不變；下回合不可防禦", "擋住並繼續"],
    ],
    [1300, 2300, 3400, 2638],
    centered_cols=(0,),
)

add_heading(doc, "6.2　狀態編碼與轉移", 2)
add_paragraphs(
    doc,
    """
令 a≥0 表示某玩家現有子彈數；令 d∈{0,1}，其中 d=1 表示本回合可以防禦，d=0 表示本回合不能防禦。把個人狀態編碼為 r=2a+d。於是 r 的整數部分同時保存兩項資訊：a=⌊r/2⌋，而 r 的奇偶就是 d。奇數狀態可防禦，偶數狀態不可防禦。

兩人的聯合狀態寫成 x=(r,s)，第一個座標屬於甲，第二個座標屬於乙。此編碼沒有遺失資訊，而且把「防禦不能連續」自動放入轉移規則。
""",
)
add_equation(doc, "r = 2a+d，　a = ⌊r/2⌋，　d = r mod 2。")
make_table(
    doc,
    ["目前 r", "L 後", "S 後", "D 後"],
    [
        ["奇數", "r+2", "r-2（r≥3）", "r-1"],
        ["正偶數", "r+3", "r-1", "不合法"],
        ["0", "3", "不合法", "不合法"],
    ],
    [1800, 2400, 2800, 2638],
    centered_cols=(0, 1, 2, 3),
)
add_labeled(doc, "例：", "r=3 代表一發子彈且可防禦。選 L 後成為 5，選 S 後成為 1，選 D 後成為 2。")

add_heading(doc, "6.3　有限期價值 V_N^z", 2)
add_paragraphs(
    doc,
    """
令 V_N^z(r,s) 表示從狀態（r,s）開始、最多再進行 N 回合的零和賽局值。若 N 回合內已分勝負，使用 +1 或 -1；若仍未結束，統一給截止收益 z。本文使用 z=-1、0、+1。

當 N=0 時不能再行動，所以 V_0^z(r,s)=z。若 N≥1，先列出本回合矩陣；立即勝負的格子填 ±1，其餘格子填相應下一狀態的 V_{N-1}^z。取此矩陣的零和賽局值，就得到 V_N^z。這是有限次倒推，因此不涉及無限極限。
""",
)
add_equation(doc, "V_0^z(r,s)=z，　　　 V_{N+1}^z = T(V_N^z)。")
add_note(doc, "索引約定", "N 一律表示「從目前狀態起尚可進行的回合數」。因此（3,3）的一回合矩陣若使用續局值 A_N、B_N，這些續局值會是剩餘 N-1 回合的數值。")

add_heading(doc, "6.4　有限矩陣賽局與可檢查證書", 2)
add_paragraphs(
    doc,
    """
設甲的混合策略為 p=(p_1,…,p_m)，其中 p_i≥0 且總和為 1。乙若選純欄 β，甲的平均收益是 Σ_αp_αM_{αβ}。甲希望選 p，使最壞一欄仍盡量高；乙則以機率 q 混合各欄，使甲任何純列的平均收益都盡量低。
""",
)
add_equation(doc, "val(M)=max_p min_β Σ_α p_αM_{αβ}=min_q max_α Σ_β M_{αβ}q_β。")
add_note(doc, "有限極小極大定理", "對任何有限實數矩陣，上式左右兩數相等，而且最大值與最小值都能取得。本文引用此基礎定理來保證每一層存在最優策略；實際數值則用下述證書核對。")
add_heading(doc, "6.4.1　上下保證證書", 3)
add_labeled(doc, "引理 1：", "若存在 p、q、v，使 p 對每一純欄的平均收益都至少 v，而 q 使每一純列的平均收益都至多 v，則 val(M)=v，且（p,q）是零和納許均衡。")
add_labeled(doc, "證明：", "甲使用 p 時，不論乙如何混合，總收益都是各純欄收益的加權平均，因此至少 v；所以甲可保證至少 v。乙使用 q 時，不論甲如何混合，總收益也是各純列收益的加權平均，因此至多 v；所以乙可保證至多 v。同一賽局不可能同時高於自己的上界或低於自己的下界，故值恰為 v。證畢。")
add_body(doc, "這個引理只用到「非負權重的加權平均仍介於各數的最小值與最大值之間」。程式求得候選 p、q 後，會逐欄與逐列檢查這些不等式。")

add_heading(doc, "6.5　Bellman 算子 T 是什麼", 2)
add_paragraphs(
    doc,
    """
給定任意未來價值函數 W，對每個狀態 x=(r,s) 建立矩陣 M_x(W)。若行動組合立即使甲勝，填 +1；若立即使甲敗，填 -1；若繼續到 y，填 W(y)。定義
""",
)
add_equation(doc, "(TW)(x)=val(M_x(W))。")
add_body(doc, "所以 T 不是一個額外策略，而是「把未來數值填入一回合矩陣，再解該矩陣」的運算。由於每位玩家最多只有 L、S、D 三種行動，對每個固定 x，M_x(W) 最多是三乘三矩陣。")

add_heading(doc, "6.5.1　單調性", 3)
add_labeled(doc, "引理 2：", "若對所有狀態都有 W(x)≤Z(x)，則 TW(x)≤TZ(x)。")
add_labeled(doc, "證明：", "M_x(W) 的立即勝負格與 M_x(Z) 相同；續局格則因 W≤Z 而逐格不大於。固定任何 p、q，較小矩陣的加權平均不大於較大矩陣的加權平均；再依序取最小值與最大值，不等號方向不變。證畢。")

add_heading(doc, "6.5.2　最大誤差不放大", 3)
add_labeled(doc, "引理 3：", "若所有與 x 有關的續局狀態都滿足 |W(y)-Z(y)|≤ε，則 |TW(x)-TZ(x)|≤ε。")
add_labeled(doc, "證明：", "兩個一回合矩陣的每一格最多相差 ε。固定 p、q 後，兩個平均收益之差是各格誤差的加權平均，絕對值最多 ε。因此，W 矩陣中任何策略的保證值，在 Z 矩陣中最多只改變 ε；交換 W、Z 再做一次，就得到雙邊差距不超過 ε。證畢。")
add_note(doc, "用途", "這個 1-Lipschitz 性質使有限個矩陣元素取極限時，矩陣賽局值也能取極限；不需要使用無限維空間的理論。")

add_heading(doc, "6.6　上下截止的單調性與極限", 2)
add_paragraphs(
    doc,
    """
定義悲觀截止與樂觀截止：
""",
)
add_equation(doc, "V̲_N=V_N^{-1}，　　　 V̄_N=V_N^{+1}。")
add_labeled(doc, "定理 1：", "對每個固定狀態 x，V̲_N(x) 隨 N 單調上升，V̄_N(x) 隨 N 單調下降，且 -1≤V̲_N(x)≤V̄_N(x)≤1。")
add_labeled(doc, "下界的起點：", "所有矩陣元素都至少 -1，所以 T(-1)≥-1，即 V̲_1≥V̲_0。")
add_labeled(doc, "下界的歸納：", "若 V̲_N≥V̲_{N-1}，由引理 2 得 T(V̲_N)≥T(V̲_{N-1})，也就是 V̲_{N+1}≥V̲_N。")
add_labeled(doc, "上界：", "同理，所有矩陣元素至多 +1，所以 T(+1)≤+1；再用引理 2 歸納，得到 V̄_{N+1}≤V̄_N。")
add_labeled(doc, "夾住：", "起始時 -1≤+1；引理 2 反覆作用給出 V̲_N≤V̄_N。所有收益都在 [-1,1]，故完整不等式成立。證畢。")
add_body(doc, "大一微積分中的有界單調數列定理說：單調上升且有上界的數列必收斂；單調下降且有下界的數列也必收斂。因此可逐狀態定義")
add_equation(doc, "V̲(x)=lim_{N→∞}V̲_N(x)，　　　 V̄(x)=lim_{N→∞}V̄_N(x)。")
add_note(doc, "目前只得到什麼？", "到這一步只證明兩條序列各自收斂，尚未證明它們收斂到同一數。無限期證明真正要完成的是 V̄(x)-V̲(x)=0。")

add_heading(doc, "6.7　極限仍滿足 Bellman 方程", 2)
add_labeled(doc, "引理 4：", "V̲=T(V̲)，且 V̄=T(V̄)。")
add_labeled(doc, "證明：", "固定狀態 x。它只有有限個行動組合，因此只會用到有限個下一狀態 y。對任意 ε>0，當 N 夠大時，這些有限個 y 都同時滿足 |V̲_N(y)-V̲(y)|≤ε。由引理 3，|T(V̲_N)(x)-T(V̲)(x)|≤ε。左邊第一項是 V̲_{N+1}(x)，令 N→∞ 即得 V̲(x)=T(V̲)(x)。上極限完全相同。證畢。")

add_heading(doc, "6.8　交換對稱與對角狀態", 2)
add_paragraphs(
    doc,
    """
把兩位玩家交換，甲的勝負會反號。記交換狀態 x*=(s,r)。由有限期矩陣逐格交換可用歸納法得到
""",
)
add_equation(doc, "V̲_N(r,s)=-V̄_N(s,r)。")
add_body(doc, "令 N→∞ 得 V̲(r,s)=-V̄(s,r)。特別地，若最後證明上下極限相同並記為 V，則 V(r,s)=-V(s,r)，所以每個對角狀態都滿足 V(r,r)=0。這包含雙方裝一發子彈且可防禦的（3,3）。")

# CHAPTER6_MIDDLE
add_heading(doc, "6.9　可由連續射擊強迫獲勝的狀態", 2)
add_paragraphs(
    doc,
    """
若乙正被射擊，為了不立刻落敗，只能選 S 或 D。設乙有 b 發子彈、目前防禦指標為 d。若 d=1，乙最多可依 D,S,D,S,… 回應 2b+1 次；若 d=0，必須先 S，再交替 D,S，最多回應 2b 次。兩種情況的最大回應次數恰好都是狀態碼 s=2b+d。

因此，若甲有至少 s+1 發子彈，甲連續射擊 s+1 次，乙一定有一次無法以 S 或 D 回應。由 a(r)=⌊r/2⌋，得到甲的充分必勝條件
""",
)
add_equation(doc, "⌊r/2⌋ ≥ s+1。")
add_body(doc, "對稱地，若 ⌊s/2⌋≥r+1，乙可強迫獲勝。把這兩類狀態的聯集記為 F。從 F 出發，勝負可在有限步內確定，所以 V̲=V̄=±1，間距為 0。")

add_heading(doc, "6.10　狀態資源的單調性", 2)
add_labeled(doc, "引理 5：", "增加甲的狀態碼不會降低 V̄；增加乙的狀態碼不會提高 V̄。相同敘述也適用於每個有限上截止值。")
add_paragraphs(
    doc,
    """
證明用剩餘回合數歸納。截止層是常數，顯然單調。假設上一層對自己的狀態碼單調，考慮 r 增加 1。

若 r 為偶數，r=2a 表示有 a 發子彈但不能防禦；r+1=2a+1 有同樣子彈且多了防禦權。r 原本的 L、S 在 r+1 都能到達相同下一狀態，並且額外多出 D，所以價值不會降低。

若 r 為奇數，r+1 表示少了目前防禦權但多了一發子彈。舊 L 可由新 L 對應，且新 L 的下一狀態多兩個狀態碼；舊 S 可由新 S 對應，新 S 的下一狀態也多兩個狀態碼。舊 D 則可由新 S 取代：面對對方 L 時，新 S 直接獲勝而舊 D 只繼續；面對 S 或 D 時，新 S 的下一狀態至少不差。依歸納假設，這些替代都不降低收益。因此 r+1 的最大保證值至少是 r 的值。

對手資源的方向可用同樣的欄比較，或先交換兩位玩家再利用反號對稱性得到。令 N→∞，有限期單調性保留到 V̄。證畢。
""",
)
add_note(doc, "為何不能只寫「顯然」？", "狀態碼加一有時是多一個防禦權，有時是多一發子彈但暫時不能防禦；兩種奇偶情形必須分開比較。")

add_heading(doc, "6.11　上下極限間距與交叉最優策略", 2)
add_paragraphs(
    doc,
    """
定義尚未消除的不確定寬度
""",
)
add_equation(doc, "g(x)=V̄(x)-V̲(x)，　　　 0≤g(x)≤2。")
add_body(doc, "由交換對稱，g(r,s)=V̄(r,s)+V̄(s,r)。我們要證明所有 x 都有 g(x)=0。")
add_paragraphs(
    doc,
    """
對每個狀態 x，從上極限矩陣 M_x(V̄) 選一個甲的最優保證策略 p_x。再到交換狀態 x*=(s,r) 選 p_{x*}。回到原狀態 x 時，讓甲使用 p_x，而乙把 p_{x*} 當作自己的欄策略。這組配法稱為交叉最優策略。它不是說兩人合作，而是利用同一個上極限問題與交換對稱，讓兩條最優性不等式能剛好相減。
""",
)

add_heading(doc, "6.11.1　交叉間距不等式的逐步證明", 3)
add_labeled(doc, "第一個保證：", "p_x 是 M_x(V̄) 的甲方最優策略，所以面對任何乙策略 q，特別是 q=p_{x*}，都有 p_x^T M_x(V̄)p_{x*}≥V̄(x)。")
add_labeled(doc, "第二個保證：", "p_{x*} 在交換狀態保證 V̄(x*)。交換矩陣的列欄並把收益反號後，正好得到 M_{x*}(V̄)_{β,α}=-M_x(V̲)_{α,β}。因此 p_x^T M_x(V̲)p_{x*}≤-V̄(x*)=V̲(x)。")
add_labeled(doc, "相減：", "用第一式減第二式，左邊的 V̄(x)-V̲(x) 就是 g(x)。右邊是同一組行動機率下兩個矩陣的平均差。若本回合立即結束，兩個矩陣都填同一個 ±1，差為 0；若繼續到 y，差就是 g(y)。")
add_equation(doc, "g(x) ≤ E_x[g(X_1)·1_{本回合未終止}]。")
add_body(doc, "E_x 表示從 x 出發，依交叉策略對有限個行動組合做加權平均；1_{條件} 在條件成立時等於 1，否則等於 0。這條式子的意思是：現在的間距只能被帶到下一個尚未結束的狀態，不會在一回合中憑空增加。")

add_heading(doc, "6.12　為何還需要資源分數", 2)
add_paragraphs(
    doc,
    """
若交叉過程必定在有限時間內進入間距已知為 0 的狀態，反覆使用上一節不等式即可得到 g=0。可是玩家能一直裝彈，所以必須排除狀態向無限大逃走。為此定義個人資源分數
""",
)
add_equation(doc, "h(r)=r+(r mod 2)=2(⌊r/2⌋+(r mod 2))。")
add_body(doc, "h 把一發子彈與一個目前可用的防禦權都算成 2 分。雙方總分 H(r,s)=h(r)+h(s)。其一回合改變如下。")
make_table(
    doc,
    ["目前奇偶", "行動", "h 的改變", "原因"],
    [
        ["奇數", "L", "+2", "多一發子彈，仍可防禦"],
        ["奇數", "S", "-2", "少一發子彈"],
        ["奇數", "D", "-2", "失去本回合防禦權"],
        ["正偶數", "L", "+4", "多一發子彈並恢復防禦權"],
        ["正偶數", "S", "0", "少一發子彈但恢復防禦權"],
    ],
    [1900, 1300, 1700, 4738],
    centered_cols=(0, 1, 2),
)
add_body(doc, "最後定義截平後的資源函數 Φ。到達終止或必勝集合 F 時令 Φ=0；其他狀態令")
add_equation(doc, "Φ(r,s)=max{10,H(r,s)}。")
add_body(doc, "因為 h 與 H 都是偶數，非終止狀態的 Φ 只可能是 10、12、14、16、…。選擇底部 10 的目的，是把所有小狀態收進有限核心，同時讓高資源區可直接用 H 的增減計算。")

add_heading(doc, "6.13　高資源區的平均不增加", 2)
add_labeled(doc, "命題 2：", "若 x 不在 F 且 H(x)≥14，則可選擇前述上極限最優策略，使交叉策略下 E[Φ(X_1)-Φ(x)]≤0。")
add_body(doc, "以下依兩個狀態碼的奇偶分三類。終止或進入 F 會把 Φ 直接降為 0，所以計算時忽略這些額外下降，只會把所得上界放鬆，不會破壞「≤0」。")

add_heading(doc, "6.13.1　偶數—偶數", 3)
add_paragraphs(
    doc,
    """
雙方都不能防禦，只能 L 或 S。對甲而言，列 S 面對欄 L 立即得到 +1，面對欄 S 則進入（r-1,s-1）；對乙而言，欄 S 面對列 L 立即給甲 -1，面對列 S 也是同一續局。因此 S/S 是一組鞍點，可選雙方都射擊。

偶數狀態射擊後 h 不變，所以 H 不變；若下一狀態剛好進入 F，Φ 反而下降。故平均改變不大於 0。
""",
)

add_heading(doc, "6.13.2　奇數—奇數", 3)
add_paragraphs(
    doc,
    """
令 ℓ_1=p_x(L)、ℓ_2=p_{x*}(L) 為兩位玩家在各自上極限最優策略中的裝彈機率。甲若面對乙純射擊 S，自己裝彈的收益是 -1，其餘行動的收益至多 +1。因此該平均收益至多 1-2ℓ_1；但 p_x 必須保證 V̄(x)，所以
""",
)
add_equation(doc, "ℓ_1≤(1-V̄(x))/2，　　　 ℓ_2≤(1-V̄(x*))/2。")
add_body(doc, "兩式相加，並用 g(x)=V̄(x)+V̄(x*)，得到 ℓ_1+ℓ_2≤1-g(x)/2。")
add_paragraphs(
    doc,
    """
奇數狀態選 L 使 h 增加 2，選 S 或 D 使 h 減少 2。因此，先把終止格當成一般轉移計算，H 的平均改變是 4(ℓ_1+ℓ_2)-4。若 L/S 或 S/L 立即終止，實際 Φ 還會多下降目前的 H。令 t 為這兩種終止組合的總機率，便有
""",
)
add_equation(doc, "E[ΔΦ]≤4(ℓ_1+ℓ_2)-4-Ht≤-2g(x)-Ht≤0。")

add_heading(doc, "6.13.3　偶數—奇數", 3)
add_paragraphs(
    doc,
    """
設甲為偶數狀態、乙為奇數狀態。乙作為欄玩家時，欄 L 被欄 S 弱支配：甲若 L，乙 S 使甲立即得 -1，不高於乙 L 的續局值；甲若 S，乙 L 使甲得 +1，而乙 S 的續局值至多 +1。因此可把乙的 L 刪除。交換狀態中同理可選一個不裝彈的奇數方最優策略。

剩下的二乘二矩陣如下，其中 b=V̄(r-1,s-2)、c=V̄(r+3,s-1)、d=V̄(r-1,s-1)。
""",
)
make_table(
    doc,
    ["甲＼乙", "S", "D"],
    [["L", "-1", "c"], ["S", "b", "d"]],
    [2400, 3619, 3619],
    centered_cols=(0, 1, 2),
)
add_paragraphs(
    doc,
    """
由狀態單調性，b≥d、c≥d。又從狀態（r+3,s-1）先射擊，面對乙 L 可立即勝，面對乙 S 則到自己的資源比 b 所在狀態更多的續局，因此 c≥b。

令甲裝彈機率為 λ。面對乙 S、D 的平均收益分別是 b-λ(1+b) 與 d+λ(c-d)。前者隨 λ 不增加，後者隨 λ 不減少；非退化時最好的 λ 在兩條直線交點：
""",
)
add_equation(doc, "λ=(b-d)/(1+b+c-d)。")
add_body(doc, "若分母退化，可選 λ=0。又因 b≤c 且 d≥-1，b-d≤1+c，故 2(b-d)≤1+b+c-d，得到 λ≤1/2。")
add_paragraphs(
    doc,
    """
令奇數方射擊機率為 σ；他不裝彈，所以 S 或 D 都使自己的 h 減少 2。偶數方 L 使 h 增加 4，S 使 h 不變。忽略終止時的額外下降，平均改變為 4λ-2；L/S 以機率 λσ 終止，而該格原本的 H 增量為 +2，所以實際還多減 H+2。於是
""",
)
add_equation(doc, "E[ΔΦ]≤4λ-2-(H+2)λσ≤0。")
add_body(doc, "奇數—偶數情形完全對稱。至此，高資源區的無限多個狀態已全部由三種奇偶情形處理。")

# CHAPTER6_CORE
add_heading(doc, "6.14　低資源有限核心與精確分數證書", 2)
add_paragraphs(
    doc,
    """
高資源證明涵蓋 H≥14。因 H 為偶數，其餘非必勝狀態都滿足 Φ=10 或 12。定義有限核心
""",
)
add_equation(doc, "C={x：x∉F 且 H(x)≤12}。")
add_body(doc, "直接依狀態碼列舉後，C 恰有 37 個有序狀態，其中 Φ=10 有 27 個，Φ=12 有 10 個。這裡的「有序」表示（r,s）與（s,r）分別計算。")
make_table(
    doc,
    ["層", "狀態數", "狀態清單"],
    [
        ["Φ=10", "27", "(0,0),(0,1),(1,0),(1,1),(1,2),(1,3),(2,1),(2,2),(2,3),(2,4),(2,5),(3,1),(3,2),(3,3),(3,4),(3,5),(3,6),(4,2),(4,3),(4,4),(4,5),(4,6),(5,2),(5,3),(5,4),(6,3),(6,4)"],
        ["Φ=12", "10", "(3,7),(4,7),(4,8),(5,5),(5,6),(6,5),(6,6),(7,3),(7,4),(8,4)"],
    ],
    [1200, 1000, 7438],
    size=8.7,
    centered_cols=(0, 1),
)

add_heading(doc, "6.14.1　為何可以包住真正的極限最優策略", 3)
add_paragraphs(
    doc,
    """
令 M_x^+=M_x(V̄_{11})，也就是把第 11 層樂觀值填進 x 的一回合矩陣。因 V̄_{11}≥V̄，矩陣 M_x^+ 的每個續局格都不小於真正上極限矩陣。另一方面 V̄(x)≥V̲_{12}(x)。

若 p_x 是真正矩陣 M_x(V̄) 的最優保證策略，它對每一欄都至少保證 V̄(x)。把矩陣元素往上放大為 M_x^+ 後，保證不會變差；再把右端從 V̄(x) 放低到 V̲_{12}(x)，不等式仍成立。因此 p_x 必位於下面集合：
""",
)
add_equation(doc, "P_x={p≥0：Σp=1，且 p^T M_x^+(:,β)≥V̲_{12}(x) 對每一欄 β}。")
add_body(doc, "P_x 是由有限條直線不等式切出的多邊形或線段。它可能比真正最優策略集合大，但一定把真正 p_x 包在裡面，所以只要對 P_x 中所有策略都能證明漂移≤0，真正策略自然也通過。")

add_heading(doc, "6.14.2　為何只檢查頂點", 3)
add_paragraphs(
    doc,
    """
固定乙策略 q 時，漂移 p^TD_xq 對 p 是一次函數。若 p 不是 P_x 的頂點，而可寫成 p=θp_1+(1-θ)p_2，其中 0<θ<1，則漂移也是兩個端點漂移的加權平均，不可能同時大於兩端點。因此最大值可移到某個頂點。固定該頂點 p 後，對 q 做同樣論證，最大值也可移到 P_{x*} 的頂點。

本題每位玩家至多三個行動，機率總和等於 1 後只剩二維。程式從「機率等於 0」與「某欄保證值恰等於下界」等邊界中選取足夠多條，解小型聯立方程，列出所有可行頂點。所有係數都是 Fraction，所以高斯消去與不等式檢查都精確。
""",
)

add_heading(doc, "6.14.3　證書結果", 3)
make_table(
    doc,
    ["檢查項目", "精確結果", "數學意義"],
    [
        ["核心狀態數", "37", "已列完所有 x∉F、H≤12 的有序狀態"],
        ["最大漂移違規數", "0", "每個頂點對皆有 p^TD_xq≤0"],
        ["Φ=10 的逃離", "4 回合內 >9/100", "任何外包策略下都有正機率離開此層"],
        ["Φ=12 的逃離", "2 回合內 >1/2", "任何外包策略下都有正機率離開此層"],
    ],
    [2300, 2200, 5138],
    centered_cols=(1,),
)
add_note(doc, "實際精確輸出", "程式顯示 core N 12 states 37 bad 0；以小數顯示的最小逃離機率約為 0.0956646187 與 0.5578499989，但 assert 比較的是分數是否分別大於 9/100 與 1/2。")

add_heading(doc, "6.15　超鞅只代表「下一步平均不增加」", 2)
add_paragraphs(
    doc,
    """
令 τ 為第一次到達立即終止或必勝集合 F 的回合。交叉策略產生隨機狀態 X_0,X_1,…；把到達 τ 後的資源分數固定為 0，寫成 Y_n=Φ(X_{min(n,τ)})。

前兩節已證明：知道目前歷史與 X_n 後，下一步 Y_{n+1} 的有限加權平均不大於 Y_n。這種非負隨機數列稱為超鞅（supermartingale）。本文不直接引用「非負超鞅收斂定理」，而把需要的部分用上下穿越方法重新證明。
""",
)

add_heading(doc, "6.15.1　上下穿越引理的初等證明", 3)
add_labeled(doc, "引理 6：", "若 Y_n≥0 且下一步條件平均不增加，則 Y_n 以機率 1 收斂。")
add_paragraphs(
    doc,
    """
固定 0≤a<b。想像一位只按規則交易的觀察者：每當 Y 降到 a 以下便買進一單位；之後第一次升到 b 以上便賣出；再重複。到第 n 步完成的「由 a 以下到 b 以上」次數記 U_n(a,b)。每次完整交易至少賺 b-a；若最後仍持有一單位，因買價至多 a、Y_n≥0，最壞只損失 a。因此路徑上總收益 G_n 滿足
""",
)
add_equation(doc, "G_n≥(b-a)U_n(a,b)-a。")
add_paragraphs(
    doc,
    """
每一步是否持有，只由以前觀察到的 Y 決定。持有時，下一步平均增量不大於 0；沒有持有時增量為 0。因此把有限步的平均增量相加，得到 E[G_n]≤0。兩式合併：
""",
)
add_equation(doc, "E[U_n(a,b)]≤a/(b-a)。")
add_paragraphs(
    doc,
    """
U_n 隨 n 只增不減，但它的平均永遠被同一常數壓住，所以發生無限多次穿越的機率只能是 0。若 Y_n 不收斂，就會有 liminf Y_n<limsup Y_n；可在兩者之間選到有理數 a<b，使序列無限多次由 a 以下跑到 b 以上，與前句矛盾。所有有理數對只有可數多個，故除去總機率 0 的例外路徑後，Y_n 必收斂。證畢。
""",
)
add_note(doc, "離散值的好處", "Y_n 只取 0、10、12、14、… 等相隔至少 2 的值。一個取離散值且收斂的數列，最後必定固定在某一個值，因為進入距離極限小於 1 的範圍後已沒有第二個可選值。")

add_heading(doc, "6.16　排除所有正的固定值", 2)
add_labeled(doc, "Φ=10：", "不論落在哪一個該層核心狀態，接下來 4 回合內離開該層的條件機率至少 9/100。因此連續 k 個四回合區塊都沒離開的機率至多 (91/100)^k，令 k→∞ 得 0。")
add_labeled(doc, "Φ=12：", "同理，每 2 回合內離開的條件機率至少 1/2，連續 k 個區塊都不離開的機率至多 (1/2)^k→0。")
add_labeled(doc, "Φ≥14：", "若兩碼皆偶數，選定的 S/S 使下一步成為奇數—奇數而 H 不變。奇數—奇數若 H 不變且未終止，只可能是 L/D 或 D/L，下一步成為一奇一偶。混合奇偶時，非終止的每個行動組合都使 H 改變 ±2；L/S 則直接終止。因此最多三步內，Φ 必改變或到達 τ，不可能永遠固定。")
add_body(doc, "Y_n 幾乎必定收斂且最後固定，但所有正固定值都已排除，所以最後只能固定在 0。也就是")
add_equation(doc, "P(τ<∞)=1。")
add_body(doc, "「以機率 1」表示例外路徑的總機率為 0；它不表示每一條形式上可能寫出的無限行動序列都不可能。")

add_heading(doc, "6.17　間距為零與無限期遊戲有值", 2)
add_paragraphs(
    doc,
    """
從交叉間距不等式出發，對未終止的下一狀態再套用同一不等式，反覆 n 次可得
""",
)
add_equation(doc, "g(x)≤E_x[g(X_n)·1_{τ>n}]≤2P_x(τ>n)。")
add_body(doc, "因 τ<∞ 的機率為 1，事件 {τ>n} 隨 n 增大縮小到空事件，故 P(τ>n)→0。右端趨近 0，而 g(x)≥0，因此 g(x)=0。")
add_labeled(doc, "主定理 3：", "對每個合法狀態 x，V̲(x)=V̄(x)。把共同值記為 V(x)，則平手截止值 V_N^0(x) 也收斂到 V(x)。")
add_labeled(doc, "證明平手截止：", "每一層起始截止值滿足 -1≤0≤+1；由 Bellman 單調性反覆作用，V̲_N≤V_N^0≤V̄_N。兩端收斂到同一個 V，夾擠定理給出 V_N^0→V。證畢。")

add_heading(doc, "6.17.1　為何 V 是真正無限期遊戲的值", 3)
add_paragraphs(
    doc,
    """
任取 ε>0。選 N 使 V̲_N(x)>V(x)-ε。甲在真正無限遊戲的前 N 回合使用悲觀截止有限賽局的最優策略：若 N 回合內結束，有限與無限收益相同；若未結束，悲觀賽局記 -1，而真正無限收益至少也是 -1。因此甲可保證期望收益大於 V-ε。

對稱地，選 N 使 V̄_N(x)<V(x)+ε。乙使用樂觀截止有限賽局的前 N 回合策略；若未結束，樂觀賽局已給甲最高可能的 +1，所以真正無限收益不會更大。乙可把甲的期望收益壓到 V+ε 以下。

因此無限期的下值與上值都等於 V。兩個有限前綴策略構成 ε-最優保證。本文沒有因此自動宣稱存在一組在所有狀態都精確達到 V 的平穩策略；那需要另一個存在性論證。
""",
)

add_heading(doc, "6.18　（3,3）第二回合策略的解析公式", 2)
add_paragraphs(
    doc,
    """
若雙方第一回合都裝彈，各有一發子彈且可以防禦，所以第二回合狀態是（3,3）。這是條件分析，並不先假設第一回合「必然」在所有均衡中都裝彈。

令 A_N=V_{N-1}^0(5,2)、B_N=V_{N-1}^0(2,1)。利用交換對稱與對角值 0，按 L、S、D 排列，（3,3）的 N 回合矩陣為：
""",
)
make_table(
    doc,
    ["甲＼乙", "L", "S", "D"],
    [
        ["L", "0", "-1", "A_N"],
        ["S", "1", "0", "-B_N"],
        ["D", "-A_N", "B_N", "0"],
    ],
    [1800, 2612, 2613, 2613],
    centered_cols=(0, 1, 2, 3),
)
add_paragraphs(
    doc,
    """
當 A_N>0、B_N>0 時，設甲混合機率為（l,s,d）。因矩陣反對稱，值為 0。甲面對乙三個純行動的平均收益依序是 s-A_Nd、-l+B_Nd、A_Nl-B_Ns。要同時保證不小於 0，前兩式給 s≥A_Nd、l≤B_Nd；第三式再給 l≥(B_N/A_N)s≥B_Nd。因此必有 l=B_Nd、s=A_Nd。再用 l+s+d=1，得到唯一策略
""",
)
add_equation(doc, "(L_N,S_N,D_N)=(B_N,A_N,1)/(1+A_N+B_N)。")
add_paragraphs(
    doc,
    """
主定理已證明 A_N→A=V(5,2)、B_N→B=V(2,1)。第 14 層的精確悲觀下界已為正，所以 A、B>0，且 A_N、B_N 最終也為正。分子分母都收斂、分母永遠至少 1，依商的極限定律可得
""",
)
add_equation(doc, "(L_N,S_N,D_N)→(B,A,1)/(1+A+B)。")
add_labeled(doc, "結論：", "有限期價值對所有狀態收斂；在（3,3）這個矩陣中，正的 A、B 又使均衡唯一，因此有限期第一步策略也收斂。這正是價值收斂能推出策略收斂所需要的額外條件。")

add_heading(doc, "6.19　計算方法與可重現性", 2)
add_paragraphs(
    doc,
    """
探索階段使用雙精度浮點數快速計算到 N=80，觀察值、間距與策略曲線。證明階段則使用 Fraction。零和矩陣的候選頂點由「總機率為 1」加上若干個緊束條件形成的小型聯立方程；高斯消去後逐一檢查所有機率非負與所有保證不等式。

核心驗證的邏輯不是抽樣：先完整列舉 37 個核心狀態，再完整列舉 P_x 與 P_{x*} 的頂點對。任何真正的上極限最優策略都在這些多面體中，而雙線性最大值必在頂點對取得，所以通過有限檢查即可涵蓋連續無限多個混合策略。
""",
)
add_code(
    doc,
    """
hi, lo = build(+1), build(-1)       # exact Fraction recursion
for x in core_states:               # all 37 ordered states
    Px  = outer_vertices(M_hi[x], lo[12, x])
    Pxs = outer_vertices(M_hi[x*], lo[12, x*])
    assert max(p.T @ Drift[x] @ q for p in Px for q in Pxs) <= 0

assert min_escape(level=10, steps=4) > Fraction(9, 100)
assert min_escape(level=12, steps=2) > Fraction(1, 2)
""",
)
add_note(doc, "軟體版本欄", "作業系統：請填寫；Python 版本：請填寫；C++ 編譯器版本：請填寫。正式投稿前請在實際執行電腦上重跑並填入。")

# CHAPTER7_START
bounds_by_n = {int(row["n"]): row for row in DATA["bounds"]}
strategy_by_h = {int(row["horizon_at_33"]): row for row in DATA["strategy"]}

add_chapter(doc, "七、結果與討論（含結論與建議）")

add_heading(doc, "7.1　主要結果總表", 2)
make_table(
    doc,
    ["研究問題", "結論", "證據層級"],
    [
        ["上下截止是否各自收斂？", "是；下界單調上升、上界單調下降", "純解析證明"],
        ["兩個極限是否相等？", "是；每個合法狀態皆 V̲=V̄", "解析證明＋37 狀態精確證書"],
        ["平手截止 V_N^0 是否收斂？", "是；由上下夾擠到共同極限 V", "純解析推論"],
        ["無限期遊戲是否有值？", "是；值為 V，雙方皆有任意 ε 的保證策略", "純解析推論"],
        ["（3,3）策略是否收斂？", "是；收斂到 (B,A,1)/(1+A+B)", "解析公式＋A、B>0"],
        ["十位小數是否已嚴格認證？", "否；目前十位數只屬浮點探索", "限制說明"],
    ],
    [2900, 3650, 3088],
)
add_note(doc, "最重要的區分", "本研究已嚴格證明「存在共同極限」；浮點計算則提供「共同極限大約是多少」。後者若要宣稱前七位小數正確，仍須額外加入向外取整區間或更深層精確分數計算。")

add_heading(doc, "7.2　無限期價值主定理", 2)
add_labeled(doc, "定理 4（總結）：", "對每個合法狀態（r,s），V_N^{-1}(r,s) 與 V_N^{+1}(r,s) 收斂到同一個 V(r,s)，而 V_N^0(r,s) 也收斂到 V(r,s)。以最終勝負 ±1、永不結束 0 為收益的無限期遊戲有值 V(r,s)。")
add_body(doc, "此結果不是從圖形外推，而是由「交叉間距不等式＋資源分數平均不增加＋有限核心精確證書＋停止後夾擠」組成。證明對所有合法狀態同時適用，並未假設子彈數有上限。")

add_heading(doc, "7.3　目前可嚴格證明的數值範圍", 2)
add_paragraphs(
    doc,
    """
使用同一套 Fraction 遞迴計算到第 14 層，可得下列包含真實極限的區間。表中端點已向外取到四位小數；內部比較仍使用完整分數。
""",
)
make_table(
    doc,
    ["量", "精確分數計算的外包區間", "用途"],
    [
        ["A=V(5,2)", "[0.7075, 0.7863]", "證明 A>0"],
        ["B=V(2,1)", "[0.6346, 0.7194]", "證明 B>0"],
        ["極限裝彈機率 L", "[0.2621, 0.2965]", "由 B/(1+A+B) 的單調性"],
        ["極限射擊機率 S", "[0.2915, 0.3248]", "由 A/(1+A+B) 的單調性"],
        ["極限防禦機率 D", "[0.3991, 0.4270]", "由 1/(1+A+B) 的單調性"],
    ],
    [2500, 3000, 4138],
    centered_cols=(1,),
)
add_paragraphs(
    doc,
    """
例如 L=B/(1+A+B) 對 B 增加、對 A 減少，所以其最小值可用 B 的下端與 A 的上端代入，最大值則反過來。S 的方向相反；D 對 A、B 都遞減。這些單調方向可由兩個分數交叉相乘直接驗證，不必用偏微分。

這組界已足以嚴格證明 A、B 為正，從而證明（3,3）矩陣的均衡唯一與策略收斂；但區間仍太寬，不能宣稱小數前七位已確定。
""",
)

add_heading(doc, "7.4　浮點數值探索", 2)
add_note(doc, "非嚴格誤差證書", "以下所有 N=80 小數均由雙精度浮點程式產生，用來顯示趨勢與提供後續區間驗證目標。它們不參與第六章的收斂證明。")

gap_rows = []
for n in (5, 10, 20, 30, 40, 50, 60, 70, 80):
    row = bounds_by_n[n]
    gap_rows.append([n, f"{row['gap52']:.10g}", f"{row['gap21']:.10g}"])
make_table(
    doc,
    ["N", "G_N(5,2)", "G_N(2,1)"],
    gap_rows,
    [1600, 4019, 4019],
    centered_cols=(0, 1, 2),
)
add_body(doc, "從 N=25 之後，每增加五回合，兩個間距的比值在這組輸出中約落在 0.302 至 0.304。這支持指數衰減猜想，但本研究的嚴格收斂證明沒有假設這個比例永遠成立。")
add_figure(
    doc,
    ROOT / "figure_gap.png",
    "圖 1　關鍵狀態的有限期上下截止間距（縱軸為對數刻度；浮點探索）",
    "折線圖顯示 N 從 1 增加到 80 時，狀態 (5,2) 與 (2,1) 的上下截止間距在對數刻度下近似直線下降。",
)

add_heading(doc, "7.4.1　第二回合策略的數值趨勢", 3)
strategy_rows = []
for h in (6, 11, 21, 31, 41, 61, 81):
    row = strategy_by_h[h]
    strategy_rows.append([h, f"{row['L']:.9f}", f"{row['S']:.9f}", f"{row['D']:.9f}"])
make_table(
    doc,
    ["（3,3）剩餘回合", "L", "S", "D"],
    strategy_rows,
    [2500, 2379, 2379, 2380],
    centered_cols=(0, 1, 2, 3),
)
add_figure(
    doc,
    ROOT / "figure_strategy.png",
    "圖 2　（3,3）有限期均衡的裝彈、射擊與防禦機率（浮點探索）",
    "折線圖顯示三個有限期均衡機率隨剩餘回合數增加而穩定，約趨近裝彈 0.2792、射擊 0.3077、防禦 0.4131。",
)
make_table(
    doc,
    ["量", "N=80 附近的浮點估計"],
    [
        ["A=V(5,2)", "0.744842400823"],
        ["B=V(2,1)", "0.675847510775"],
        ["L", "0.2791962356"],
        ["S", "0.3076983951"],
        ["D", "0.4131053693"],
    ],
    [4700, 4938],
    centered_cols=(1,),
)

add_heading(doc, "7.5　兩個容易混淆的零", 2)
add_labeled(doc, "間距的零：", "本研究已證明 g(3,3)=V̄(3,3)-V̲(3,3)=0；事實上所有合法狀態的 g 都是 0。這表示上下截止對無限未來的歧見完全消失。")
add_labeled(doc, "價值的零：", "共同價值又因交換對稱滿足 V(3,3)=-V(3,3)，所以 V(3,3)=0。這表示在對稱狀態中，公平值是 0。兩個敘述的理由不同：前者來自收斂主證明，後者來自對稱性。")

add_heading(doc, "7.6　對研究問題的回答", 2)
for item in (
    "有限 N 可以精確倒推：每一層只需求有限個至多三乘三矩陣。",
    "上下截止各自收斂：這是有界單調數列的直接結果。",
    "上下極限相等：交叉策略把間距推到未來，資源分數與有限證書保證過程以機率 1 到達間距為 0 的狀態。",
    "平手截止收斂且極限是無限期值：由上下夾擠與有限前綴 ε-保證策略得到。",
    "條件在雙方先裝彈後，（3,3）的有限期第一步策略收斂，極限由 A、B 的解析公式決定。",
    "目前沒有嚴格七位小數認證；已有完整存在性證明、精確粗界與可作目標的高精度浮點估計。",
):
    add_numbered(doc, item)

add_heading(doc, "7.7　討論與限制", 2)
add_paragraphs(
    doc,
    """
本研究最初試圖直接找 G_N 的指數上界。最後的證明改從「過程必停」下手，雖然沒有給出漂亮的全域收斂速率，卻已足以排除任何非零振盪並證明真正極限。這顯示在無限狀態問題中，先找合適的資源量往往比直接追逐小數速度更有效。

證明中有一小段是電腦輔助的，但其範圍明確：只有 37 個有序狀態，所有輸入是分數，所有候選來自有限個頂點，最後條件是精確不等式。解析部分負責把無限問題縮成這個有限證書；程式並沒有替代整個證明。

目前的主要限制是數值精度認證。N=80 的上下差在浮點輸出中約為 10^{-8}，但浮點運算本身沒有向外取整，不能據此說前七位小數必然正確。下一步可用區間算術：每次加減乘除都把下端向下、上端向上取整，並讓矩陣值也輸出包含真值的區間。

此外，本文證明無限期遊戲有值與 ε-最優策略，並證明（3,3）有限期第一步策略收斂；尚未完整分類每個狀態是否都有唯一的精確平穩均衡。若遊戲規則改變，例如允許連續防禦或同時射擊也結束，現有資源分數與必勝條件也需重新分析。
""",
)

add_heading(doc, "7.8　結論", 2)
add_paragraphs(
    doc,
    """
本研究完成一個無限狀態、無固定期限的同時行動遊戲之價值證明。有限期悲觀值與樂觀值分別單調收斂；交叉最優策略給出間距不等式；資源分數 Φ 在高資源區可手算證明平均不增加，低資源區則縮成 37 個分數狀態的有限證書。上下極限因此相等，平手截止值收斂，而共同極限就是無限期遊戲的值。

在雙方第一回合皆裝彈的條件下，第二回合狀態（3,3）的有限期混合策略也收斂。若 A=V(5,2)、B=V(2,1)，極限為（B,A,1）/(1+A+B)。浮點探索估計約為裝彈 0.2791962356、射擊 0.3076983951、防禦 0.4131053693；目前嚴格證明的是收斂本身與較寬的分數區間，而不是這些小數的全部位數。
""",
)

add_heading(doc, "7.9　後續建議", 2)
for item in (
    "完成向外取整區間版本，正式認證 A、B 與三個策略機率的前七位小數。",
    "嘗試從資源過程導出 P(τ>N) 或 G_N 的明確收斂速率；不必預設一定是 0.31 的五回合倍率。",
    "研究所有狀態的極限一回合矩陣是否唯一，進一步判斷有限期策略能否逐狀態收斂。",
    "探討是否存在可在所有狀態同時精確達值的平穩策略。",
    "修改防禦冷卻、同時射擊規則或初始子彈數，比較證明中哪些結構仍保留。",
):
    add_bullet(doc, item)


add_chapter(doc, "八、參考文獻")
add_body(doc, "正文中的理論背景與軟體說明依下列資料整理。網頁查閱日期：2026 年 8 月 22 日。", first_indent=False)
references = [
    "von Neumann, J. (1928). Zur Theorie der Gesellschaftsspiele. Mathematische Annalen, 100, 295–320. https://doi.org/10.1007/BF01448847",
    "Nash, J. F., Jr. (1950). Equilibrium points in n-person games. Proceedings of the National Academy of Sciences, 36(1), 48–49. https://doi.org/10.1073/pnas.36.1.48",
    "Shapley, L. S. (1953). Stochastic games. Proceedings of the National Academy of Sciences, 39(10), 1095–1100. https://doi.org/10.1073/pnas.39.10.1095",
    "Osborne, M. J., & Rubinstein, A. (1994). A Course in Game Theory. MIT Press. https://mitpress.mit.edu/9780262650403/a-course-in-game-theory/",
    "von Neumann, J., & Morgenstern, O. (1944). Theory of Games and Economic Behavior. Princeton University Press. https://cart.press.princeton.edu/theory-of-games-and-economic-behavior-pb-395.html",
    "Python Software Foundation. fractions—Rational numbers. Python 3 documentation. https://docs.python.org/3/library/fractions.html",
    "Python Software Foundation. Floating-point arithmetic: Issues and limitations. Python 3 tutorial. https://docs.python.org/3/tutorial/floatingpoint.html",
]
for i, ref in enumerate(references, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{i}. {ref}")
    set_run_font(r, size=10.5)


add_chapter(doc, "九、附錄")
add_heading(doc, "附錄 A　主要符號表", 2)
make_table(
    doc,
    ["符號", "定義", "用途"],
    [
        ["L、S、D", "裝彈、射擊、防禦", "純行動"],
        ["r=2a+d", "個人狀態碼", "同時保存子彈與防禦權"],
        ["x=(r,s)", "甲、乙的聯合狀態", "Bellman 遞迴的索引"],
        ["V_N^z", "剩餘 N 回合、截止收益 z 的值", "有限期模型"],
        ["T", "把未來值填入矩陣再取其賽局值", "Bellman 算子"],
        ["V̲_N、V̄_N", "截止 -1、+1 的有限期值", "無限期上下界"],
        ["V̲、V̄", "上下截止的逐狀態極限", "共同值證明前的兩個候選"],
        ["g", "V̄-V̲", "要證明為 0 的間距"],
        ["p_x", "M_x(V̄) 的甲方最優策略", "交叉策略"],
        ["F", "可由連續射擊強迫勝負的狀態集合", "停止集合"],
        ["h、H", "個人及雙方資源分數", "控制狀態不逃向無限"],
        ["Φ", "終止為 0，否則 max{10,H}", "非負超鞅"],
        ["τ", "第一次到達終止或 F 的時間", "停止性與間距夾擠"],
        ["A、B", "V(5,2)、V(2,1)", "（3,3）極限策略公式"],
    ],
    [1600, 4200, 3838],
    size=9.2,
)

add_heading(doc, "附錄 B　完整證明依賴表", 2)
make_table(
    doc,
    ["順序", "命題", "只需要的數學", "用途"],
    [
        ["1", "有限矩陣有值", "極小極大定理；證書用加權平均", "定義 T"],
        ["2", "T 單調且誤差不放大", "有限和、最大最小", "傳遞上下界與極限"],
        ["3", "V̲_N、V̄_N 單調收斂", "數學歸納、有界單調數列", "定義 V̲、V̄"],
        ["4", "極限固定點與交換對稱", "有限狀態極限、矩陣轉置", "選 p_x 並寫 g"],
        ["5", "交叉間距不等式", "兩個最優性不等式相減", "把 g 推向下一步"],
        ["6", "高資源漂移≤0", "奇偶分類、二乘二代數", "處理無限多狀態"],
        ["7", "低資源漂移≤0", "Fraction、聯立方程、頂點", "完成所有狀態"],
        ["8", "Y_n 收斂且只能到 0", "上下穿越、幾何級數", "證明 τ<∞ 幾乎必然"],
        ["9", "g=0", "期望上界與夾擠", "得到共同極限"],
        ["10", "無限遊戲有值", "有限前綴 ε-策略", "連接有限與無限"],
        ["11", "（3,3）策略收斂", "三元一次不等式、商的極限", "第二回合答案"],
    ],
    [700, 2800, 3100, 3038],
    size=8.9,
    centered_cols=(0,),
)

add_heading(doc, "附錄 C　精確核心驗證輸出", 2)
add_code(
    doc,
    """
core N 12 states 37 bad 0
level 10 steps 4 min_escape 0.09566461871352176
CERTIFIED level 10, 4 steps: min_escape > 9/100
level 12 steps 2 min_escape 0.5578499989272458
CERTIFIED level 12, 2 steps: min_escape > 1/2
""",
)
add_body(doc, "上列小數只為閱讀方便；程式實際 assert 比較 Fraction。完整程式碼公開位置：請填寫。若以紙本繳交，建議同時附上原始碼與執行說明的電子檔。")

add_heading(doc, "附錄 D　演算法摘要", 2)
add_code(
    doc,
    """
actions(r):
    A = [L]
    if r >= 2: add S
    if r is odd: add D

next(r,L) = r+2 if r odd else r+3
next(r,S) = r-2 if r odd else r-1
next(r,D) = r-1

V(0,r,s) = cutoff
V(n,r,s) = value of the matrix whose entry is
             +1 for S/L, -1 for L/S,
             V(n-1,next(r,a),next(s,b)) otherwise
""",
)
add_body(doc, "驗證用指令與檔名：請填寫。建議至少記錄成功輸出、執行時間與程式雜湊值，讓評審能確認所附程式就是產生報告結果的版本。")

add_heading(doc, "附錄 E　研究紀錄表", 2)
make_table(
    doc,
    ["日期", "版本", "修改或實驗內容", "結果／問題", "下一步"],
    [
        ["請填寫", "請填寫", "建立有限 N 模型", "請填寫", "請填寫"],
        ["請填寫", "請填寫", "觀察五回合間距倍率", "請填寫", "請填寫"],
        ["請填寫", "請填寫", "建立上下截止證明", "請填寫", "請填寫"],
        ["請填寫", "請填寫", "完成核心 Fraction 驗證", "請填寫", "請填寫"],
        ["請填寫", "請填寫", "完成區間精度認證", "請填寫", "請填寫"],
    ],
    [1300, 1100, 3000, 2200, 2038],
    size=8.9,
    centered_cols=(0, 1),
)

add_heading(doc, "附錄 F　投稿前檢查表", 2)
for item in (
    "封面所有『請填寫』欄位已完成。",
    "目錄頁碼已在最終版重新核對。",
    "所有圖、表在正文都有解釋，圖上浮點結果皆標明非嚴格證書。",
    "核心驗證已在投稿使用的電腦重新執行，37 個狀態、bad 0 與兩個 escape assert 均通過。",
    "參考文獻格式已依主辦單位規定統一。",
    "若主張七位小數，已附向外取整區間，且不是只附四捨五入的小數。",
):
    add_bullet(doc, "□ " + item)


# Update fields on open and normalize all sections.
settings = doc.settings._element
old = settings.find(qn("w:updateFields"))
if old is not None:
    settings.remove(old)
node = OxmlElement("w:updateFields")
node.set(qn("w:val"), "true")
settings.append(node)

for section in doc.sections:
    setup_section(section)
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""

doc.save(OUT)
print(OUT)
